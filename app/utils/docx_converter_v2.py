"""
DOCX to PDF Converter with Exact Page Structure Preservation
Converts DOCX files to PDF maintaining exact spacing, page breaks, and layout

Required libraries:
pip install python-docx reportlab pillow
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.units import inch, cm, mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 Table, TableStyle, Image, PageBreak,
                                 ListFlowable, ListItem, KeepTogether)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
from PIL import Image as PILImage
import os
import re


def get_paragraph_alignment(para):
    """Get paragraph alignment"""
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: TA_LEFT,
        WD_ALIGN_PARAGRAPH.CENTER: TA_CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT: TA_RIGHT,
        WD_ALIGN_PARAGRAPH.JUSTIFY: TA_JUSTIFY,
    }
    return alignment_map.get(para.alignment, TA_LEFT)


def get_text_with_formatting(paragraph):
    """Extract text with inline formatting preserved - DO NOT TRIM"""
    parts = []
    for run in paragraph.runs:
        text = run.text  # Keep exact text with all spaces
        if text == '':
            continue
       
        # Escape special characters for ReportLab
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
       
        # Apply formatting
        if run.bold and run.italic:
            text = f'<b><i>{text}</i></b>'
        elif run.bold:
            text = f'<b>{text}</b>'
        elif run.italic:
            text = f'<i>{text}</i>'
       
        if run.underline:
            text = f'<u>{text}</u>'
           
        parts.append(text)
   
    return ''.join(parts)  # Return exact text without stripping


def pts_to_reportlab(pt_value):
    """Convert Word points to ReportLab points"""
    if pt_value is None:
        return None
    return float(pt_value)


def get_spacing_before(para):
    """Get space before paragraph in points"""
    if para.paragraph_format.space_before:
        return pts_to_reportlab(para.paragraph_format.space_before.pt)
    return 0


def get_spacing_after(para):
    """Get space after paragraph in points"""
    if para.paragraph_format.space_after:
        return pts_to_reportlab(para.paragraph_format.space_after.pt)
    return 0


def get_line_spacing(para):
    """Get line spacing"""
    if para.paragraph_format.line_spacing:
        # Line spacing in Word can be a multiplier or exact points
        line_spacing = para.paragraph_format.line_spacing
        if isinstance(line_spacing, float):
            return line_spacing
        elif hasattr(line_spacing, 'pt'):
            return line_spacing.pt
    return None


def get_heading_style(para, styles):
    """Determine if paragraph is a heading and return appropriate style"""
    style_name = para.style.name if para.style else ''
   
    # Get font size
    font_size = 11
    if para.runs and para.runs[0].font.size:
        font_size = para.runs[0].font.size.pt
   
    # Get spacing
    space_before = get_spacing_before(para)
    space_after = get_spacing_after(para)
    line_spacing_value = get_line_spacing(para)
   
    # Calculate leading
    leading = font_size * 1.2 if line_spacing_value is None else font_size * line_spacing_value
   
    if 'Heading 1' in style_name or 'Title' in style_name:
        return ParagraphStyle(
            'Heading1',
            parent=styles['Heading1'],
            fontSize=font_size,
            textColor=colors.HexColor('#2E75B6'),
            spaceAfter=space_after,
            spaceBefore=space_before,
            leading=leading,
            alignment=get_paragraph_alignment(para)
        )
    elif 'Heading 2' in style_name:
        return ParagraphStyle(
            'Heading2',
            parent=styles['Heading2'],
            fontSize=font_size,
            textColor=colors.HexColor('#2E75B6'),
            spaceAfter=space_after,
            spaceBefore=space_before,
            leading=leading,
            alignment=get_paragraph_alignment(para)
        )
    elif 'Heading 3' in style_name or 'Heading' in style_name:
        return ParagraphStyle(
            'Heading3',
            parent=styles['Heading3'],
            fontSize=font_size,
            textColor=colors.HexColor('#2E75B6'),
            spaceAfter=space_after,
            spaceBefore=space_before,
            bold=True,
            leading=leading,
            alignment=get_paragraph_alignment(para)
        )
   
    return None


def get_paragraph_style(para, styles):
    """Get or create appropriate style for paragraph with EXACT spacing"""
    # Check if it's a heading
    heading_style = get_heading_style(para, styles)
    if heading_style:
        return heading_style
   
    # Get alignment
    alignment = get_paragraph_alignment(para)
   
    # Get font properties from first run
    font_size = 11
    is_bold = False
    is_italic = False
   
    if para.runs:
        first_run = para.runs[0]
        if first_run.font.size:
            font_size = first_run.font.size.pt
        is_bold = first_run.bold if first_run.bold is not None else False
        is_italic = first_run.italic if first_run.italic is not None else False
   
    # Get exact spacing from Word
    space_before = get_spacing_before(para)
    space_after = get_spacing_after(para)
    line_spacing_value = get_line_spacing(para)
   
    # Calculate leading (line height)
    if line_spacing_value:
        leading = font_size * line_spacing_value
    else:
        leading = font_size * 1.2  # Default Word line spacing
   
    # Get indentation
    left_indent = 0
    first_line_indent = 0
    if para.paragraph_format.left_indent:
        left_indent = pts_to_reportlab(para.paragraph_format.left_indent.pt)
    if para.paragraph_format.first_line_indent:
        first_line_indent = pts_to_reportlab(para.paragraph_format.first_line_indent.pt)
   
    # Create custom style with EXACT spacing
    style = ParagraphStyle(
        'CustomParagraph',
        parent=styles['Normal'],
        fontSize=font_size,
        alignment=alignment,
        spaceAfter=space_after,
        spaceBefore=space_before,
        leading=leading,
        leftIndent=left_indent,
        firstLineIndent=first_line_indent
    )
   
    return style


def extract_images_from_doc(doc):
    """Extract all images from document with their relationships"""
    images = {}
    image_index = 0
   
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_data = rel.target_part.blob
                images[rel.rId] = {
                    'data': image_data,
                    'index': image_index
                }
                image_index += 1
            except:
                pass
   
    return images


def get_images_in_paragraph(para, doc_images):
    """Get images embedded in a paragraph"""
    para_images = []
   
    # Check for inline shapes
    for run in para.runs:
        if 'graphic' in run._element.xml:
            # Extract relationship ID
            for rel_id in re.findall(r'r:embed="([^"]+)"', run._element.xml):
                if rel_id in doc_images:
                    para_images.append(doc_images[rel_id]['data'])
   
    return para_images


def is_list_paragraph(para):
    """Check if paragraph is part of a list"""
    if para._element.pPr is not None:
        numPr = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
        if numPr is not None:
            return True
    return False


def has_page_break_before(para):
    """Check if paragraph has page break before"""
    if para._element.pPr is not None:
        pageBreakBefore = para._element.pPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pageBreakBefore')
        if pageBreakBefore is not None:
            return True
    return False


def convert_table_to_reportlab(docx_table):
    """Convert DOCX table to ReportLab table with styling"""
    data = []
   
    for row in docx_table.rows:
        row_data = []
        for cell in row.cells:
            # Get cell text with basic formatting
            cell_text = ''
            for para in cell.paragraphs:
                text = get_text_with_formatting(para)
                if text:  # Don't use strip() to preserve spaces
                    cell_text += text + '<br/>'
            row_data.append(cell_text.rstrip('<br/>'))
        data.append(row_data)
   
    if not data:
        return None
   
    # Calculate column widths based on content
    num_cols = len(data[0]) if data else 0
    if num_cols == 0:
        return None
   
    available_width = 6.5 * inch
    col_width = available_width / num_cols
   
    # Create table
    table = Table(data, colWidths=[col_width] * num_cols)
   
    # Apply styling
    style = TableStyle([
        # Header row styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
       
        # Body styling
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.black),
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('TOPPADDING', (0, 1), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
       
        # Grid
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ])
   
    table.setStyle(style)
    return table


def docx_to_pdf(docx_path, pdf_path):
    """
    Convert DOCX to PDF while preserving EXACT structure and spacing
   
    Args:
        docx_path (str): Path to input DOCX file
        pdf_path (str): Path to output PDF file
    """
   
    # Load document
    doc = Document(docx_path)
   
    # Get page settings from Word document
    section = doc.sections[0]
    page_width = section.page_width.inches if section.page_width else 8.27  # A4 width
    page_height = section.page_height.inches if section.page_height else 11.69  # A4 height
    left_margin = section.left_margin.inches if section.left_margin else 1
    right_margin = section.right_margin.inches if section.right_margin else 1
    top_margin = section.top_margin.inches if section.top_margin else 1
    bottom_margin = section.bottom_margin.inches if section.bottom_margin else 1
   
    # Extract all images
    doc_images = extract_images_from_doc(doc)
   
    # Create PDF with EXACT page settings from Word
    pdf_doc = SimpleDocTemplate(
        pdf_path,
        pagesize=(page_width*inch, page_height*inch),
        rightMargin=right_margin*inch,
        leftMargin=left_margin*inch,
        topMargin=top_margin*inch,
        bottomMargin=bottom_margin*inch
    )
   
    # Story container
    story = []
    styles = getSampleStyleSheet()
   
    # Track list items
    list_items = []
    in_list = False
   
    # Process each paragraph and table
    for block in doc.element.body:
        # Handle tables
        if block.tag.endswith('tbl'):
            # Add any pending list items
            if list_items:
                story.append(ListFlowable(list_items, bulletType='bullet', start='•'))
                list_items = []
                in_list = False
           
            # Find the table object
            for table in doc.tables:
                if table._element == block:
                    rl_table = convert_table_to_reportlab(table)
                    if rl_table:
                        story.append(rl_table)
                    break
       
        # Handle paragraphs
        elif block.tag.endswith('p'):
            # Find the paragraph object
            para = None
            for p in doc.paragraphs:
                if p._element == block:
                    para = p
                    break
           
            if para is None:
                continue
           
            # Check for page break before
            if has_page_break_before(para):
                if list_items:
                    story.append(ListFlowable(list_items, bulletType='bullet', start='•'))
                    list_items = []
                    in_list = False
                story.append(PageBreak())
           
            # Check for images in paragraph
            para_images = get_images_in_paragraph(para, doc_images)
           
            # Add images
            for img_data in para_images:
                # Add any pending list items before image
                if list_items:
                    story.append(ListFlowable(list_items, bulletType='bullet', start='•'))
                    list_items = []
                    in_list = False
               
                try:
                    img = PILImage.open(BytesIO(img_data))
                   
                    # Calculate dimensions - preserve original size as much as possible
                    aspect = img.height / float(img.width)
                   
                    # Get available width and height - use 100% to match DOCX exactly
                    available_width = (page_width - left_margin - right_margin) * inch
                    available_height = (page_height - top_margin - bottom_margin) * inch  # Use 100% of available height
                   
                    # Scale to fit width
                    img_width = min(img.width, available_width)
                    img_height = img_width * aspect
                   
                    # If height is too large, scale down based on height
                    if img_height > available_height:
                        img_height = available_height
                        img_width = img_height / aspect
                   
                    # Create image
                    img_buffer = BytesIO()
                    img.save(img_buffer, format='PNG')
                    img_buffer.seek(0)
                   
                    rl_image = Image(img_buffer, width=img_width, height=img_height)
                    story.append(rl_image)
                except Exception as e:
                    print(f"Error processing image: {e}")
           
            # Get paragraph text - DO NOT STRIP OR TRIM
            text = get_text_with_formatting(para)
           
            # Always process paragraph, even if empty (to preserve spacing)
            if text == '':
                # Empty paragraph - add exact spacing from Word
                space_after = get_spacing_after(para)
                if space_after > 0:
                    story.append(Spacer(1, space_after))
                else:
                    # Add minimal spacer for empty line
                    font_size = 11
                    if para.runs and para.runs[0].font.size:
                        font_size = para.runs[0].font.size.pt
                    story.append(Spacer(1, font_size * 1.2))
                continue
           
            # Check if this is a list item
            if is_list_paragraph(para):
                # Create list item
                item_style = get_paragraph_style(para, styles)
                list_items.append(ListItem(Paragraph(text, item_style), leftIndent=20))
                in_list = True
            else:
                # Add any pending list items
                if list_items:
                    story.append(ListFlowable(list_items, bulletType='bullet', start='•'))
                    list_items = []
                    in_list = False
               
                # Get appropriate style with EXACT spacing
                para_style = get_paragraph_style(para, styles)
               
                # Add paragraph with exact spacing
                story.append(Paragraph(text, para_style))
   
    # Add any remaining list items
    if list_items:
        story.append(ListFlowable(list_items, bulletType='bullet', start='•'))
   
    # Build PDF
    try:
        pdf_doc.build(story)
        print(f"Successfully converted {docx_path} to {pdf_path}")
        print(f"Page settings: {page_width:.2f}\" x {page_height:.2f}\"")
        print(f"Margins: L:{left_margin:.2f}\" R:{right_margin:.2f}\" T:{top_margin:.2f}\" B:{bottom_margin:.2f}\"")
        return True
    except Exception as e:
        print(f"Error building PDF: {e}")
        import traceback
        traceback.print_exc()
        return False


# # Example usage
# if __name__ == "__main__":
#     input_docx = "test_docx1.docx"  # Your input file
#     output_pdf = "output_exact.pdf"  # Output file
   
#     if not os.path.exists(input_docx):
#         print(f"Error: Input file '{input_docx}' not found!")
#     else:
#         success = docx_to_pdf(input_docx, output_pdf)
       
#         if success:
#             print(f"\nConversion complete! PDF saved to: {output_pdf}")
#             print("The PDF maintains exact spacing and page structure from the DOCX file.")
#         else:
#             print("\nConversion failed. Check error messages above.")
